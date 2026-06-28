from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "manual"
LOGO = ROOT / "app" / "static" / "img" / "logo-gt.png"
VERSION = "2.1.1"
GOLD = "D6A619"
INK = "2D3033"
MUTED = "747981"


CONTENT = {
    "pt": {
        "filename": "GT-Stock-Manager-Manual-PT.docx",
        "title": "Manual de Utilização",
        "subtitle": "GT Stock Manager",
        "edition": "Versão 2.1.1 | Português",
        "contents": "Índice",
        "intro": (
            "Este manual descreve a utilização operacional e administrativa do GT Stock Manager. "
            "Os ecrãs e permissões apresentados dependem do perfil atribuído a cada utilizador."
        ),
        "chapters": [
            ("1. Acesso e idioma", "login", [
                "Abra o endereço do sistema, introduza o nome de utilizador e a senha e selecione Entrar.",
                "Use o seletor PT/EN no cabeçalho. A preferência fica associada à conta e permanece após terminar e voltar a iniciar sessão.",
                "No primeiro acesso, altere a senha quando o sistema solicitar. Nunca partilhe credenciais.",
            ]),
            ("2. Dashboard", "dashboard", [
                "O Dashboard resume produtos, ruturas, alertas, movimentos, requisições e processos de Procurement.",
                "Os cartões e gráficos são atualizados a partir dos movimentos e estados registados no sistema.",
                "Clique nos atalhos e tabelas para abrir o relatório ou processo correspondente.",
            ]),
            ("3. Produtos e economato", "products", [
                "Consulte, filtre e edite produtos. O código é gerado pelo sistema; o nome original nunca é traduzido automaticamente.",
                "Pode indicar um nome opcional em inglês para apresentação bilingue sem alterar o dado original.",
                "Defina unidade, preço unitário, stock mínimo e se o produto deve ser monitorizado. Itens inativos ou não monitorizados não geram alertas de reposição.",
                "A quantidade existente só deve ser corrigida por utilizadores autorizados, com motivo obrigatório e movimento de auditoria.",
            ]),
            ("4. Movimentos de stock", "movement", [
                "Escolha Entrada, Saída, Devolução ou Acerto. Os campos mudam conforme a ação.",
                "Entradas exigem origem ou fornecedor; saídas exigem departamento e responsável.",
                "O sistema bloqueia quantidades inválidas e qualquer operação que produza saldo negativo.",
                "Acertos manuais exigem justificação e ficam registados na auditoria.",
            ]),
            ("5. Requisições SR e aprovações", "requisition", [
                "Selecione o departamento, o gestor operacional em serviço e os produtos com stock disponível.",
                "O valor total é calculado pelo preço unitário e determina o perfil aprovador através da matriz.",
                "O aprovador pode aprovar tudo, rejeitar tudo ou aprovar parcialmente. Cada quantidade rejeitada exige uma observação.",
                "A aprovação não reduz stock. A saída ocorre uma única vez quando o Gestor de Stock confirma a emissão/entrega.",
                "Use Cancelar ou Voltar para sair do fluxo sem criar movimentos.",
            ]),
            ("6. Procurement / Non-Stock", "procurement", [
                "Crie o pedido com tipo de necessidade, prioridade, orçamento estimado, descrição, requisitos técnicos e HSE.",
                "O fluxo segue: necessidade, TdR, aprovações HOD e Diretor do Terminal, orçamento, classificação, concurso, avaliação técnica e financeira, adjudicação, PO, receção e arquivo.",
                "A rota final de aprovação é calculada pela matriz usando o valor atualizado do processo.",
                "O TdR e o formulário podem ser visualizados, descarregados em PDF e enviados por e-mail.",
            ]),
            ("7. Reposição de stock", "replenishment", [
                "O Gestor de Stock seleciona produtos monitorizados, informa quantidade e preço estimado e cria o pedido de compra.",
                "A seleção pode incluir itens sugeridos pelo stock mínimo e outros produtos ativos.",
                "Após aprovação e PO, registe apenas quantidades fisicamente recebidas. Cada receção cria uma entrada ligada ao processo.",
            ]),
            ("8. Relatórios", "reports", [
                "Estão disponíveis relatórios de stock, movimentos, requisições, Procurement/reposições e itens que requerem atenção.",
                "Aplique filtros antes de exportar. CSV, Excel e PDF usam o idioma selecionado e preservam os dados introduzidos pelos utilizadores.",
                "Produtos inativos ou sem monitorização não aparecem no relatório de atenção.",
            ]),
            ("9. Utilizadores", "users", [
                "Administradores podem criar e editar contas, definir perfil, departamento, idioma e canais de notificação.",
                "O departamento é opcional; campos inválidos apresentam uma mensagem clara em vez de erro técnico.",
                "Remover acesso desativa a conta e preserva o histórico. O SuperAdmin e a própria conta têm proteções adicionais.",
            ]),
            ("10. Perfis e permissões", "profiles", [
                "Perfis representam funções e controlam os módulos e ações disponíveis. Podem ser criados, editados e removidos quando não estão em uso.",
                "A matriz de aprovação referencia perfis, não categorias de produtos.",
                "Aplique o princípio do menor privilégio: atribua apenas os acessos necessários para a responsabilidade do utilizador.",
            ]),
            ("11. Configurações", "settings", [
                "Gira categorias, departamentos e a matriz de aprovações. A remoção preserva registos históricos através de desativação quando necessário.",
                "A matriz define intervalos de valor, modalidade e perfil aprovador. Evite intervalos incompletos ou sobrepostos.",
                "O reset total de stock exige permissão, confirmação textual e código de segurança; cria movimentos de acerto e não apaga o histórico.",
            ]),
            ("12. Notificações, documentos e importações", ["notifications", "documents", "imports"], [
                "As notificações abrem diretamente a requisição ou processo relacionado e são criadas no idioma do destinatário.",
                "E-mail e WhatsApp seguem a preferência e configuração do utilizador. Documentos de stock permanecem ligados ao respetivo movimento ou produto.",
                "Na importação, reveja a pré-visualização e o relatório de erros antes de confirmar. A gravação é transacional e auditável.",
            ]),
            ("13. Auditoria e segurança", "audit", [
                "A auditoria apresenta utilizador, ação, módulo, registo, valores anterior/novo e IP/dispositivo no idioma selecionado.",
                "Valores internos permanecem canónicos na base de dados e são traduzidos apenas na apresentação.",
                "Reveja regularmente contas ativas, perfis, matriz, stock negativo, falhas de importação e processos pendentes.",
            ]),
            ("14. Utilização móvel", "mobile-requisition", [
                "Em telemóveis, o menu fica recolhido no botão superior e os formulários/tabelas adaptam-se à largura do ecrã.",
                "Aprovações, notificações e consultas podem ser realizadas sem deslocamento horizontal.",
                "Para tarefas extensas de configuração ou importação, recomenda-se um ecrã maior.",
            ]),
        ],
        "roles_title": "Responsabilidades por perfil",
        "roles": [
            ("SuperAdmin", "Configuração total, utilizadores, perfis, matriz, auditoria e ações protegidas."),
            ("Administrador", "Operação administrativa ampla, sem alterar proteções exclusivas do SuperAdmin."),
            ("Gestor de Stock", "Movimentos, emissão/entrega, reposição, receção e ajustes autorizados."),
            ("Chefe do Terminal", "Aprovação de requisições e TdR conforme perfil e matriz."),
            ("Gestor Operacional", "Criação de pedidos e aprovação HOD quando configurado."),
            ("Procurement / Comités", "Classificação, concurso, avaliações, PO, receção e arquivo conforme permissões."),
            ("Utilizador", "Criação e acompanhamento dos próprios pedidos permitidos."),
        ],
        "support": "Em caso de erro, registe o número do processo, a página, a hora e a mensagem apresentada. Não altere dados para contornar validações.",
    },
    "en": {
        "filename": "GT-Stock-Manager-User-Manual-EN.docx",
        "title": "User Manual",
        "subtitle": "GT Stock Manager",
        "edition": "Version 2.1.1 | English",
        "contents": "Contents",
        "intro": (
            "This manual explains the operational and administrative use of GT Stock Manager. "
            "Available screens and actions depend on the profile assigned to each user."
        ),
        "chapters": [
            ("1. Access and language", "login", [
                "Open the system address, enter your username and password, and select Sign in.",
                "Use the PT/EN selector in the header. The preference is stored on the account and remains after signing out and signing in again.",
                "On first access, change the password when prompted. Never share credentials.",
            ]),
            ("2. Dashboard", "dashboard", [
                "The Dashboard summarises products, stock-outs, warnings, movements, requisitions and Procurement processes.",
                "Cards and charts update from recorded movements and workflow statuses.",
                "Use the shortcuts and tables to open the related report or process.",
            ]),
            ("3. Products and store", "products", [
                "View, filter and edit products. The system generates the code; the original name is never translated automatically.",
                "An optional English display name can be entered without replacing the original business data.",
                "Define unit, unit price, minimum stock and monitoring. Inactive or unmonitored items do not create replenishment warnings.",
                "Existing quantity adjustments require permission, a mandatory reason and an audit movement.",
            ]),
            ("4. Stock movements", "movement", [
                "Select Entry, Issue, Return or Adjustment. Relevant fields change with the action.",
                "Entries require an origin or supplier; issues require a destination department and responsible person.",
                "The system blocks invalid quantities and any operation that would create negative stock.",
                "Manual adjustments require a reason and remain in the audit trail.",
            ]),
            ("5. SR requisitions and approvals", "requisition", [
                "Select the department, Operational Manager on duty and products with available stock.",
                "The total value is calculated from unit prices and determines the approver profile through the matrix.",
                "The reviewer may approve all, reject all or partially approve. Every rejected quantity requires an observation.",
                "Approval does not reduce stock. The Stock Manager records the issue/delivery once.",
                "Use Cancel or Back to leave the workflow without creating movements.",
            ]),
            ("6. Procurement / Non-stock", "procurement", [
                "Create a request with need type, priority, estimated budget, description, technical requirements and HSE requirements.",
                "The workflow is: need, ToR, HOD and Terminal Manager approvals, budget, classification, tender, technical and financial evaluations, award, PO, receipt and archive.",
                "The final approval route is calculated from the matrix using the current process value.",
                "The ToR and request form can be previewed, downloaded as PDF and sent by email.",
            ]),
            ("7. Stock replenishment", "replenishment", [
                "The Stock Manager selects monitored products, enters quantities and estimated prices, and creates the purchase request.",
                "The selection may include minimum-stock suggestions and other active products.",
                "After approval and PO, record only physically received quantities. Each receipt creates a stock entry linked to the process.",
            ]),
            ("8. Reports", "reports", [
                "Reports cover stock, movements, requisitions, Procurement/replenishments and items requiring attention.",
                "Apply filters before exporting. CSV, Excel and PDF follow the selected language and preserve user-entered data.",
                "Inactive or unmonitored products are excluded from the attention report.",
            ]),
            ("9. Users", "users", [
                "Administrators can create and edit accounts and assign profile, department, language and notification channels.",
                "Department is optional; invalid fields show a clear message instead of a technical error.",
                "Removing access deactivates the account and preserves history. SuperAdmin and self-removal have additional protection.",
            ]),
            ("10. Profiles and permissions", "profiles", [
                "Profiles represent functions and control available modules and actions. They can be created, edited and removed when unused.",
                "The approval matrix references profiles, not product categories.",
                "Apply least privilege: assign only the access required for the user's responsibilities.",
            ]),
            ("11. Settings", "settings", [
                "Manage categories, departments and the approval matrix. Removal preserves historical records through deactivation where required.",
                "The matrix defines value ranges, procurement method and approver profile. Avoid incomplete or overlapping ranges.",
                "Full stock reset requires permission, typed confirmation and a security code; it creates adjustment movements and keeps history.",
            ]),
            ("12. Notifications, documents and imports", ["notifications", "documents", "imports"], [
                "Notifications open the related requisition or process directly and are created in the recipient's language.",
                "Email and WhatsApp follow the user's preference and configuration. Stock documents remain linked to the related movement or product.",
                "For imports, review the preview and error report before confirmation. Saving is transactional and auditable.",
            ]),
            ("13. Audit and security", "audit", [
                "The audit trail shows user, action, module, record, previous/new values and IP/device in the selected language.",
                "Internal values remain canonical in the database and are translated only for display.",
                "Regularly review active accounts, profiles, matrix rules, negative stock, import failures and pending processes.",
            ]),
            ("14. Mobile use", "mobile-requisition", [
                "On phones, the menu is collapsed into the top button and forms/tables adapt to screen width.",
                "Approvals, notifications and enquiries can be completed without horizontal scrolling.",
                "A larger screen is recommended for extensive settings or import tasks.",
            ]),
        ],
        "roles_title": "Profile responsibilities",
        "roles": [
            ("SuperAdmin", "Full configuration, users, profiles, matrix, audit and protected actions."),
            ("Administrator", "Broad administrative operation without SuperAdmin-only protections."),
            ("Stock Manager", "Movements, issue/delivery, replenishment, receipt and authorised adjustments."),
            ("Terminal Manager", "Requisition and ToR approval according to profile and matrix."),
            ("Operational Manager", "Request creation and HOD approval where configured."),
            ("Procurement / Committees", "Classification, tender, evaluations, PO, receipt and archive according to permissions."),
            ("User", "Creation and tracking of permitted own requests."),
        ],
        "support": "When reporting an error, record the process number, page, time and displayed message. Do not change data to bypass validations.",
    },
}


def shade(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    fill = OxmlElement("w:shd")
    fill.set(qn("w:fill"), color)
    tc_pr.append(fill)


def set_cell_text(cell, text, bold=False, color=INK):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def configure_document(document, title):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, before, after in [
        ("Title", 25, 0, 12),
        ("Heading 1", 16, 16, 8),
        ("Heading 2", 13, 12, 6),
        ("Heading 3", 11, 8, 4),
    ]:
        style = document.styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(INK)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "GT Caption" not in document.styles:
        caption = document.styles.add_style("GT Caption", WD_STYLE_TYPE.PARAGRAPH)
        caption.font.name = "Aptos"
        caption.font.size = Pt(8)
        caption.font.italic = True
        caption.font.color.rgb = RGBColor.from_string(MUTED)
        caption.paragraph_format.space_after = Pt(8)

    header = section.header
    table = header.add_table(rows=1, cols=2, width=Inches(6.85))
    table.columns[0].width = Inches(1.35)
    table.columns[1].width = Inches(5.5)
    if LOGO.exists():
        table.cell(0, 0).paragraphs[0].add_run().add_picture(str(LOGO), width=Inches(1.05))
    p = table.cell(0, 1).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"{title} | v{VERSION}")
    run.bold = True
    run.font.name = "Aptos"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Gestão de Terminais, SA | Departamento de Tecnologia e Informação | by Layton Taimo | ")
    run.font.name = "Aptos"
    run.font.size = Pt(7.5)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    add_page_field(p)


def add_cover(document, data):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(55)
    if LOGO.exists():
        p.add_run().add_picture(str(LOGO), width=Inches(2.6))
    p = document.add_paragraph(data["title"], style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph(data["subtitle"])
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor.from_string(GOLD)
    p = document.add_paragraph(data["edition"])
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p = document.add_paragraph("Gestão de Terminais, SA")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(85)
    p = document.add_paragraph("by Layton Taimo")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.color.rgb = RGBColor.from_string(INK)
    document.add_page_break()


def add_contents(document, data):
    document.add_heading(data["contents"], level=1)
    document.add_paragraph(data["intro"])
    for title, _image, _steps in data["chapters"]:
        document.add_paragraph(title)
    document.add_paragraph(f"15. {data['roles_title']}")
    document.add_page_break()


def add_screenshot(document, language, image_name, caption):
    if not image_name:
        return
    image = OUT / "screenshots" / language / f"{image_name}.png"
    if not image.exists():
        return
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together = True
    width = Inches(2.8) if image_name.startswith("mobile-") else Inches(6.55)
    p.add_run().add_picture(str(image), width=width)
    caption_p = document.add_paragraph(caption, style="GT Caption")
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_chapter(document, language, title, image_name, steps):
    document.add_heading(title, level=1)
    for index, step in enumerate(steps, 1):
        p = document.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        prefix = p.add_run(f"{index}. ")
        prefix.bold = True
        p.add_run(step)
    image_names = image_name if isinstance(image_name, list) else [image_name]
    for item in image_names:
        add_screenshot(document, language, item, f"{title} - {item}" if len(image_names) > 1 else title)
    document.add_page_break()


def add_roles(document, data):
    document.add_heading(data["roles_title"], level=1)
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(1.75)
    table.columns[1].width = Inches(4.9)
    headers = ("Perfil", "Responsabilidade") if data["filename"].endswith("PT.docx") else ("Profile", "Responsibility")
    for index, value in enumerate(headers):
        shade(table.rows[0].cells[index], GOLD)
        set_cell_text(table.rows[0].cells[index], value, bold=True)
    for role, responsibility in data["roles"]:
        cells = table.add_row().cells
        set_cell_text(cells[0], role, bold=True)
        set_cell_text(cells[1], responsibility)
    document.add_paragraph()
    document.add_heading("Suporte e boas práticas" if data["filename"].endswith("PT.docx") else "Support and good practice", level=2)
    document.add_paragraph(data["support"])
    document.add_paragraph(
        "Documento controlado | GT Stock Manager 2.1.1 | by Layton Taimo"
        if data["filename"].endswith("PT.docx")
        else "Controlled document | GT Stock Manager 2.1.1 | by Layton Taimo"
    )


def build(language):
    data = CONTENT[language]
    document = Document()
    configure_document(document, data["title"])
    document.core_properties.title = f"{data['title']} - GT Stock Manager"
    document.core_properties.author = "Layton Taimo"
    document.core_properties.subject = "GT Stock Manager 2.1.1"
    add_cover(document, data)
    add_contents(document, data)
    for chapter in data["chapters"]:
        add_chapter(document, language, *chapter)
    add_roles(document, data)
    output = OUT / data["filename"]
    document.save(output)
    print(output)


if __name__ == "__main__":
    OUT.mkdir(parents=True, exist_ok=True)
    build("pt")
    build("en")
